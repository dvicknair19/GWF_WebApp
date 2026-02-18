import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuration for Template
TEMPLATE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) + '/templates'
# Try to find the correct template file dynamically or use specific name
TEMPLATE_FILENAME = "TEMPLATE 2026 GWFMOA_(Client)_(Vendor)_(mo,yr).docx"

def _sanitize_filename(name: str) -> str:
    # Remove invalid filename characters
    sanitized = re.sub(r'[<>:"/\\|?*]', '', name)
    # Replace multiple spaces/underscores with single underscore
    sanitized = re.sub(r'[\s_]+', '_', sanitized)
    return sanitized.strip(' _')

def get_template_path():
    # First check for the specific new template
    path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILENAME)
    if os.path.exists(path):
        return path
    
    # Fallback: look for any docx in templates folder that looks right
    if os.path.exists(TEMPLATE_DIR):
        for f in os.listdir(TEMPLATE_DIR):
            if f.endswith('.docx') and 'GWFMOA' in f:
                return os.path.join(TEMPLATE_DIR, f)
    
    # Final fallback to env var or default
    return os.getenv('TEMPLATE_PATH', path)

def generate_vendor_profile(client_name, vendor_name, research_data):
    template_path = get_template_path()
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")

    doc = Document(template_path)

    # Populate document using the logic ported from document_service.py
    # Populate document using the logic ported from document_service.py
    _populate_document(doc, research_data, client_name)

    # Legacy Placeholder Replacement (Safety net for simple templates)
    # Replaces [CLIENT_NAME] etc in logic
    for p in doc.paragraphs:
        if '[CLIENT_NAME]' in p.text:
            p.text = p.text.replace('[CLIENT_NAME]', client_name)
        if '[VENDOR_NAME]' in p.text:
            p.text = p.text.replace('[VENDOR_NAME]', vendor_name)
    
    # Save to temporary file
    output_filename = f"Generated_{client_name}_{vendor_name}.docx"
    safe_filename = _sanitize_filename(output_filename)
    output_path = os.path.join('/tmp', safe_filename)
    doc.save(output_path)
    
    return output_path

# --- Helper functions ported and adapted from document_service.py ---

def _populate_document(doc, vendor_data, client_name=None):
    """Populate document tables and paragraphs with vendor data."""
    # 1. Search paragraphs for "Report Created For:"
    if client_name:
        for p in doc.paragraphs:
            if "report created for" in p.text.lower():
                # Append client name to the end of the paragraph
                run = p.add_run(f" {client_name}")
                # Copy font style from the first run if it exists
                if p.runs:
                    first_run = p.runs[0]
                    run.bold = first_run.bold
                    run.italic = first_run.italic
                    run.underline = first_run.underline
                    run.font.name = first_run.font.name
                    run.font.size = first_run.font.size
                    run.font.color.rgb = first_run.font.color.rgb
                    if first_run.style:
                         run.style = first_run.style

    # 2. Populate tables
    for table in doc.tables:
        _populate_table(table, vendor_data)

def _populate_table(table, vendor_data):
    """Populate a single table with vendor data."""
    # Special handlers
    _handle_vendor_profile_paragraph(table, vendor_data)
    _handle_vendor_reseller(table, vendor_data)

    for row in table.rows:
        if len(row.cells) >= 2:
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            label_text = label_cell.text.strip().lower()

            # Mapping logic
            if "company type" in label_text:
                _set_cell_value(value_cell, vendor_data.get("company_type", ""))
            elif "fiscal year end" in label_text:
                _set_cell_value(value_cell, vendor_data.get("fiscal_year_end", ""))
            elif "estimated annual revenue" in label_text or "revenue" in label_text:
                _set_cell_value(value_cell, vendor_data.get("estimated_annual_revenue", ""))
            elif "employees" in label_text or "employee" in label_text:
                _set_cell_value(value_cell, vendor_data.get("employees", ""))
            elif "competitors" in label_text and "core" in label_text:
                competitors = vendor_data.get("competitors_core", [])
                if isinstance(competitors, list):
                    _set_cell_value(value_cell, ", ".join(competitors))
                else:
                    _set_cell_value(value_cell, str(competitors) if competitors else "")
            elif "recent news" in label_text or "news" in label_text:
                news_list = vendor_data.get("recent_news", [])
                _set_news_list(value_cell, news_list)

def _handle_vendor_reseller(table, vendor_data):
    """Handle 'Vendor / Reseller' table header mapping."""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if "vendor / reseller" in cell.text.lower():
                # Target cell is the one directly below (next row, same column)
                if i + 1 < len(table.rows):
                    target_row = table.rows[i + 1]
                    if j < len(target_row.cells):
                        target_cell = target_row.cells[j]
                        # Prioritize matched name, fallback to generic name in data
                        name = vendor_data.get("matched_vendor_name") or vendor_data.get("vendor_name", "")
                        _set_cell_value(target_cell, name)
                return # Stop after first match

def _handle_vendor_profile_paragraph(table, vendor_data):
    """Handle special vendor profile paragraph placement."""
    for i, row in enumerate(table.rows):
        if len(row.cells) >= 1:
            label_text = row.cells[0].text.strip().lower()
            # Look for 'Vendor Profile' header and populate the cell below it
            if "vendor profile" in label_text and i + 1 < len(table.rows):
                next_row = table.rows[i + 1]
                if len(next_row.cells) >= 2:
                    target_cell = next_row.cells[1]
                    vendor_profile = vendor_data.get("vendor_profile_paragraph", "")
                    _set_italic_text(target_cell, vendor_profile)
                break

def _set_cell_value(cell, value):
    """Set cell text value safely."""
    # Clear existing
    for p in cell.paragraphs:
        p.clear()
    # Add new
    if not cell.paragraphs:
        cell.add_paragraph()
    
    cell.paragraphs[0].text = str(value) if value else "N/A"

def _set_italic_text(cell, text):
    """Set text with italic formatting."""
    for p in cell.paragraphs:
        p.clear()
    
    if not cell.paragraphs:
        cell.add_paragraph()
        
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text) if text else "N/A")
    run.italic = True

def _add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def _set_news_list(cell, news_list):
    """Populate news items as bullet points with hyperlinks."""
    # Remove all existing paragraph elements from the cell's XML
    tc = cell._tc
    for p_elem in tc.findall(qn('w:p')):
        tc.remove(p_elem)
    # Reset: add exactly one empty paragraph (python-docx requires at least one)
    cell.add_paragraph()

    if not news_list:
        cell.paragraphs[0].text = "No recent news available"
        return

    if not isinstance(news_list, list):
        news_list = [news_list]

    for item in news_list:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        # Add bullet character as the first run
        bullet_run = p.add_run("• ")
        bullet_run.font.size = Pt(10)
        if isinstance(item, dict):
            title = item.get("title", item.get("headline", ""))
            url = item.get("url", "")
            if url and url.startswith("http"):
                _add_hyperlink(p, title, url)
            else:
                p.add_run(title)
        else:
            p.add_run(str(item))

    # Remove the initial placeholder paragraph now that news paragraphs exist
    tc.remove(tc.findall(qn('w:p'))[0])
